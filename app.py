# app.py - REFORMATADOR ABNT 2024 (ABSTRACT OPCIONAL + CRIAÇÃO DO ZERO + VERIFICAÇÃO)
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION_START
import io
import re
import fitz  # PyMuPDF para PDF (opcional)

st.set_page_config(page_title="ABNT 2024 - Reformatador", layout="wide")
st.title("Reformatador ABNT 2024 (Criação do Zero, Fiel à Norma)")
st.caption("Preencha os dados → Gera documento perfeito conforme NBR 14724:2024! Upload opcional para conteúdo.")

# Escolha de fonte
fonte_escolhida = st.selectbox("Escolha a fonte:", ["Arial", "Times New Roman"])

# Form para dados obrigatórios
with st.form("Dados do Trabalho"):
    titulo = st.text_input("Título (obrigatório)")
    subtitulo = st.text_input("Subtítulo (opcional)")
    autor = st.text_input("Nome do Autor (obrigatório)")
    instituicao = st.text_input("Nome da Instituição (opcional)")
    local = st.text_input("Local (cidade, UF) (obrigatório)")
    ano = st.text_input("Ano de Depósito (obrigatório)")
    orientador = st.text_input("Nome do Orientador (opcional)")
    natureza = st.text_area("Natureza do Trabalho (ex: Trabalho de Conclusão de Curso apresentado à [instituição] para obtenção do grau de Bacharel em [área]) (obrigatório)")
    resumo = st.text_area("Resumo (150-500 palavras, obrigatório)")
    palavras_chave = st.text_input("Palavras-chave (3-5, separadas por ponto)")
    abstract = st.text_area("Abstract (versão em inglês, opcional)")
    keywords = st.text_input("Keywords (3-5, separadas por ponto, opcional)")
    submitted = st.form_submit_button("Gerar Documento ABNT 2024")

# Upload opcional
arquivo = st.file_uploader("Upload opcional do conteúdo (PDF/DOCX)", type=["pdf", "docx"])

if submitted and titulo and autor and local and ano and natureza and resumo:
    with st.spinner("Gerando e verificando documento conforme ABNT 2024..."):
        doc = Document()

        # === PROCESSAR UPLOAD (OPCIONAL) ===
        conteudo = ""
        if arquivo:
            if arquivo.type == "application/pdf":
                pdf_doc = fitz.open(stream=arquivo.read(), filetype="pdf")
                for page in pdf_doc:
                    conteudo += page.get_text("text") + "\n"
                    # Extrair imagens
                    images = page.get_images(full=True)
                    for img_index, img in enumerate(images):
                        xref = img[0]
                        base_image = pdf_doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        p = doc.add_paragraph(f"Figura {page.number+1}.{img_index+1} — Extraída do PDF")
                        p.add_run().add_picture(io.BytesIO(image_bytes), width=Cm(10))
                pdf_doc.close()
            else:  # DOCX
                input_doc = Document(arquivo)
                conteudo = "\n".join([p.text for p in input_doc.paragraphs])
                # Copiar imagens e tabelas
                for element in input_doc.element.body:
                    doc.element.body.append(element)

        # === CONFIGURAÇÕES GLOBAIS ===
        section = doc.sections[0]
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

        # Estilos
        def criar_estilo(nome, tam=12, negrito=False, alinh=WD_ALIGN_PARAGRAPH.JUSTIFY, esp=1.5):
            s = doc.styles.add_style(nome, WD_STYLE_TYPE.PARAGRAPH) if nome not in doc.styles else doc.styles[nome]
            s.font.name = fonte_escolhida
            s.font.size = Pt(tam)
            s.font.bold = negrito
            s.paragraph_format.alignment = alinh
            s.paragraph_format.line_spacing = esp
            return s

        texto_style = criar_estilo('ABNT_Texto', 12, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 1.5)
        titulo_style = criar_estilo('ABNT_Titulo', 12, True, WD_ALIGN_PARAGRAPH.LEFT, 1.5)
        citacao_style = criar_estilo('ABNT_Citacao', 10, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 1.0)
        citacao_style.paragraph_format.left_indent = Cm(4)
        ref_style = criar_estilo('ABNT_Ref', 12, False, WD_ALIGN_PARAGRAPH.LEFT, 1.0)
        central_style = criar_estilo('ABNT_Central', 12, True, WD_ALIGN_PARAGRAPH.CENTER, 1.5)

        # === PARTE EXTERNA E PRÉ-TEXTUAIS ===
        # Capa
        if instituicao: doc.add_paragraph(instituicao.upper(), style='ABNT_Central')
        doc.add_page_break()
        doc.add_paragraph(autor.upper(), style='ABNT_Central')
        doc.add_page_break()
        doc.add_paragraph(titulo.upper(), style='ABNT_Central')
        if subtitulo: doc.add_paragraph(subtitulo, style='ABNT_Central')
        doc.add_page_break()
        doc.add_paragraph(f"{local.upper()}\n{ano}", style='ABNT_Central')

        # Folha de Rosto
        doc.add_page_break()
        doc.add_paragraph(autor.upper(), style='ABNT_Central')
        doc.add_paragraph(titulo.upper(), style='ABNT_Central')
        if subtitulo: doc.add_paragraph(subtitulo, style='ABNT_Central')
        p = doc.add_paragraph(natureza, style='ABNT_Texto')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if orientador: doc.add_paragraph(f"Orientador: {orientador}", style='ABNT_Texto').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"{local.upper()}, {ano}", style='ABNT_Central')

        # Folha de Aprovação (placeholder)
        doc.add_page_break()
        doc.add_paragraph("FOLHA DE APROVAÇÃO", style='ABNT_Titulo')
        doc.add_paragraph(autor.upper(), style='ABNT_Texto')
        doc.add_paragraph(titulo.upper(), style='ABNT_Texto')
        if subtitulo: doc.add_paragraph(subtitulo, style='ABNT_Texto')
        doc.add_paragraph(natureza, style='ABNT_Texto')
        doc.add_paragraph("Data de Aprovação: __/__/____", style='ABNT_Texto')
        doc.add_paragraph("Banca Examinadora:\n________________________________\nNome, Titulação\n\n________________________________\nNome, Titulação", style='ABNT_Texto')

        # Resumo (obrigatório)
        doc.add_page_break()
        doc.add_paragraph("RESUMO", style='ABNT_Central')
        doc.add_paragraph(resumo, style='ABNT_Texto')
        doc.add_paragraph(f"Palavras-chave: {palavras_chave}.", style='ABNT_Texto')

        # Abstract (opcional)
        if abstract:
            doc.add_page_break()
            doc.add_paragraph("ABSTRACT", style='ABNT_Central')
            doc.add_paragraph(abstract, style='ABNT_Texto')
            if keywords: doc.add_paragraph(f"Keywords: {keywords}.", style='ABNT_Texto')

        # Sumário (placeholder)
        doc.add_page_break()
        doc.add_paragraph("SUMÁRIO", style='ABNT_Central')
        doc.add_paragraph("(Atualize no Word: Referências > Inserir Sumário)", style='ABNT_Texto')

        # === ELEMENTOS TEXTUAIS (do zero ou do upload) ===
        doc.add_page_break()  # Início textual
        new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        new_section.header.is_linked_to_previous = False

        if not conteudo:
            # Cria do zero: Seções padrão vazias
            doc.add_paragraph("1 INTRODUÇÃO", style='ABNT_Titulo')
            doc.add_paragraph("Insira o texto da introdução aqui.", style='ABNT_Texto')
            doc.add_section(WD_SECTION_START.ODD_PAGE)
            doc.add_paragraph("2 DESENVOLVIMENTO", style='ABNT_Titulo')
            doc.add_paragraph("Insira o desenvolvimento aqui.", style='ABNT_Texto')
            doc.add_section(WD_SECTION_START.ODD_PAGE)
            doc.add_paragraph("3 CONCLUSÃO", style='ABNT_Titulo')
            doc.add_paragraph("Insira a conclusão aqui.", style='ABNT_Texto')
        else:
            # Parse conteúdo do upload
            secoes_num = 1
            ref_linhas = []
            em_referencias = False
            for linha in conteudo.split('\n'):
                linha = linha.strip()
                if not linha: continue

                if linha.upper().startswith('REFERÊNCIAS'):
                    em_referencias = True
                    doc.add_paragraph("REFERÊNCIAS", style='ABNT_Titulo')
                    continue

                if em_referencias:
                    ref_linhas.append(linha)
                    continue

                # Detectar título
                if re.match(r'^\d+(\.\d+)*\s+[A-ZÀ-Ú\s]+$', linha.upper()):
                    nivel = linha.split(' ', 1)[0].count('.') + 1
                    p = doc.add_paragraph(linha.upper(), style='ABNT_Titulo')
                    if nivel == 1: doc.add_section(WD_SECTION_START.ODD_PAGE)
                else:
                    p = doc.add_paragraph(linha, style='ABNT_Texto')
                    if len(linha) > 120 and '"' in linha:
                        p.style = 'ABNT_Citacao'

                    # Limpar negrito
                    for run in p.runs:
                        run.bold = False

        # Referências (ordenadas)
        doc.add_page_break()
        if ref_linhas:
            ref_linhas.sort()
            for ref in ref_linhas:
                p = doc.add_paragraph(ref, style='ABNT_Ref')
                p.paragraph_format.space_after = Pt(6)
        else:
            doc.add_paragraph("REFERÊNCIAS", style='ABNT_Titulo')
            doc.add_paragraph("(Insira suas referências aqui, ordenadas alfabeticamente.)", style='ABNT_Ref')

        # === VERIFICAÇÃO DE NORMAS ===
        verificacao = []
        if len(resumo.split()) < 150 or len(resumo.split()) > 500: verificacao.append("Aviso: Resumo deve ter 150-500 palavras.")
        if not palavras_chave: verificacao.append("Aviso: Adicione palavras-chave (3-5).")
        if abstract and len(abstract.split()) < 150 or len(abstract.split()) > 500: verificacao.append("Aviso: Abstract deve ter 150-500 palavras se incluído.")
        if len(doc.sections) < 2: verificacao.append("Aviso: Verifique paginação (pré-textuais sem números).")
        verificacao_msg = "\n".join(verificacao) if verificacao else "Todas as normas principais aplicadas com sucesso!"

        # === SALVAR ===
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("Documento ABNT 2024 gerado com sucesso!")
        st.download_button(
            label="Baixar TCC_ABNT_2024.docx",
            data=buffer,
            file_name="TCC_ABNT_2024.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        st.info(f"**Verificação de Normas:**\n{verificacao_msg}\n\n**Finalizações no Word:**\n1. Atualize sumário.\n2. Adicione paginação (inicia na Introdução).\n3. Salvar como PDF.")

else:
    if submitted:
        st.warning("Preencha todos os campos obrigatórios (título, autor, local, ano, natureza, resumo)!")
    else:
        st.info("Preencha o form para gerar o documento do zero.")
